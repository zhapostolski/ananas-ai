"""
Generate Ananas AI Platform Master Document (Word .docx)
"""

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ORANGE = RGBColor(0xFE, 0x50, 0x00)
DARK = RGBColor(0x1A, 0x1A, 0x2E)
GRAY = RGBColor(0x64, 0x64, 0x64)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT = RGBColor(0xF9, 0xF9, 0xF9)
BG_HEADER = RGBColor(0xFE, 0x50, 0x00)
BG_ALT = RGBColor(0xFF, 0xF0, 0xEB)

doc = Document()

# ── Page margins ───────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2.2)
section.bottom_margin = Cm(2.2)

# ── Default paragraph style ────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)
style.font.color.rgb = DARK


# ── Helpers ────────────────────────────────────────────────────────────────
def set_cell_bg(cell, color: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    hex_color = f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), val.get("val", "single"))
            el.set(qn("w:sz"), val.get("sz", "4"))
            el.set(qn("w:space"), val.get("space", "0"))
            el.set(qn("w:color"), val.get("color", "auto"))
            tcBorders.append(el)
    tcPr.append(tcBorders)


def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = ORANGE
    # Bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    hex_c = f"{ORANGE[0]:02X}{ORANGE[1]:02X}{ORANGE[2]:02X}"
    bottom.set(qn("w:color"), hex_c)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = DARK
    return p


def body(text, space_after=4, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = DARK
    return p


def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent = Cm(0.8)
    if bold_prefix:
        rb = p.add_run(bold_prefix)
        rb.bold = True
        rb.font.size = Pt(10.5)
        rb.font.color.rgb = DARK
        r = p.add_run(text)
        r.font.size = Pt(10.5)
        r.font.color.rgb = DARK
    else:
        r = p.add_run(text)
        r.font.size = Pt(10.5)
        r.font.color.rgb = DARK
    return p


def spacer(n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)


def styled_table(headers, rows, col_widths=None):
    """Creates a clean table with orange header row."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        set_cell_bg(cell, BG_HEADER)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = WHITE

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.add_row()
        bg = BG_ALT if ri % 2 == 0 else WHITE
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            if isinstance(val, tuple):
                # (text, bold, color)
                text, bold, color = val
                run = p.add_run(text)
                run.bold = bold
                run.font.size = Pt(9.5)
                if color:
                    run.font.color.rgb = color
                else:
                    run.font.color.rgb = DARK
            else:
                run = p.add_run(str(val))
                run.font.size = Pt(9.5)
                run.font.color.rgb = DARK

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    return table


def info_box(title, lines, color=BG_ALT):
    """Shaded info callout box."""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, color)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    if title:
        r = p.add_run(title + "\n")
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = ORANGE
    for line in lines:
        r2 = p.add_run(line + "\n")
        r2.font.size = Pt(10)
        r2.font.color.rgb = DARK
    return table


# ═══════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
p.paragraph_format.space_after = Pt(6)
run = p.add_run("ANANAS AI PLATFORM")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = ORANGE

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(4)
run2 = p2.add_run("Intelligence Infrastructure for Marketing and Business Operations")
run2.font.size = Pt(13)
run2.font.color.rgb = DARK

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_after = Pt(40)
run3 = p3.add_run("Architecture | Project Plan | Budget | Roadmap")
run3.font.size = Pt(11)
run3.font.color.rgb = GRAY
run3.italic = True

# Divider
p_div = doc.add_paragraph()
p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr = p_div._p.get_or_add_pPr()
pBdr = OxmlElement("w:pBdr")
bot = OxmlElement("w:bottom")
bot.set(qn("w:val"), "single")
bot.set(qn("w:sz"), "12")
bot.set(qn("w:space"), "4")
bot.set(qn("w:color"), f"{ORANGE[0]:02X}{ORANGE[1]:02X}{ORANGE[2]:02X}")
pBdr.append(bot)
pPr.append(pBdr)

spacer(2)

meta_data = [
    ("Prepared for", "Denis Boskoski, Ananas.mk"),
    ("Prepared by", "AI Platform Team"),
    ("Version", "1.0"),
    ("Date", "March 2026"),
    ("Classification", "Internal — Confidential"),
]
for label, value in meta_data:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    rl = p.add_run(f"{label}:  ")
    rl.bold = True
    rl.font.size = Pt(10.5)
    rl.font.color.rgb = GRAY
    rv = p.add_run(value)
    rv.font.size = Pt(10.5)
    rv.font.color.rgb = DARK

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════
# SECTION 1 — EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════════════════════

heading1("1. Executive Summary")
body(
    "Ananas AI Platform is a purpose-built intelligence infrastructure for Ananas.mk, "
    "North Macedonia's largest e-commerce marketplace. It transforms raw marketing and "
    "business data into structured daily intelligence — delivered automatically to the "
    "right people at the right time, every morning before the team starts work."
)
spacer()
body(
    "The platform is built on three core components: a fleet of specialist AI agents that "
    "run each morning and analyse every major marketing channel in depth; a structured data "
    "layer that stores all outputs, metrics, and history; and two distribution interfaces "
    "that make intelligence accessible without asking for it."
)
spacer()
info_box(
    "What the platform delivers",
    [
        "- Daily performance intelligence across all paid channels, CRM, reputation, and marketing operations",
        "- A live portal at ai.ananas.mk for deep analysis, historical review, and team collaboration",
        "- Automatic morning briefs to Microsoft Teams channels and executive email by 07:30 daily",
        "- An AI chat interface with full Ananas business context available to all authorised users",
        "- A foundation that expands from marketing into commercial, finance, and operations",
    ],
)
spacer()
body(
    "Phase 1 is operational. Five specialist agents run daily on AWS EC2, the portal is live, "
    "GA4 is integrated and returning real data, and the first users have logged in. "
    "The platform now moves into Phase 2: expanding agent coverage, completing integrations, "
    "and rolling out to the full team."
)

spacer()
heading2("Key numbers")
styled_table(
    ["Metric", "Value"],
    [
        ("Live GA4 data (March 2026)", "~464k sessions/month, ~215k users, ~€13.4M revenue"),
        ("Agents in Phase 1", "5 specialist agents, daily cron 06:00-07:30"),
        ("Agents in Phase 2", "11 additional agents across 11 intelligence domains"),
        ("Total agent coverage (P1+P2)", "16 agents covering marketing, commercial, and ops"),
        ("Phase 1 infrastructure cost", "~€150-220/month (AWS EC2)"),
        ("Phase 2 infrastructure cost", "~€300-450/month (expanded),"),
        ("Total 18-month platform cost", "~€55,000-85,000 (all-in: infra + AI + build)"),
    ],
    col_widths=[6, 10],
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════
# SECTION 2 — THE IDEA
# ═══════════════════════════════════════════════════════════════════════════

heading1("2. The Idea")
body(
    "Marketing teams at fast-growing e-commerce companies face the same structural problem: "
    "the data exists, but getting it into the right hands, in the right format, at the right "
    "time, requires constant manual effort. Reports are built manually. Insights arrive late "
    "or not at all. Decisions are made without full context."
)
spacer()
body(
    "The Ananas AI Platform solves this by making intelligence automatic. Instead of people "
    "pulling data, the platform pushes structured intelligence to them. Every morning, before "
    "the team starts work, five specialist agents have already analysed every major channel, "
    "identified anomalies, and delivered a brief to Teams and email."
)
spacer()
body(
    "This is not a chatbot, a reporting tool, or a dashboard replacement. It is a new layer "
    "of infrastructure that sits between raw data sources and the team — continuously "
    "processing, structuring, and surfacing what matters."
)

spacer()
heading2("Why now")
styled_table(
    ["Business situation", "What the platform does about it"],
    [
        (
            "No Google Shopping campaigns despite 250,000+ products",
            "Performance Agent tracks Shopping Impression Share daily and escalates when zero campaigns are detected",
        ),
        (
            "Trustpilot rating at 2.0 — profile not yet claimed",
            "Reputation Agent monitors review count, sentiment, and response rate with critical-level alerts",
        ),
        (
            "Sales driven by marketing-budget coupons — masking true acquisition efficiency",
            "Marketing Ops Agent tracks coupon dependency ratio and contribution margin waterfall",
        ),
        (
            "No email lifecycle automations — cart recovery, churn flows not live",
            "CRM & Lifecycle Agent reports 0% recovery rate daily until automations are activated",
        ),
        (
            "8-person team managing GA4, 4 Google Ads accounts, 20 Meta campaigns, CRM, Trustpilot",
            "Cross-Channel Brief synthesises all channels into a single morning brief for the team and CEO",
        ),
        (
            "No centralised institutional memory — campaigns, learnings, and decisions lost over time",
            "Knowledge Retrieval Agent (Phase 2) surfaces past campaigns and decisions from Confluence and Jira",
        ),
    ],
    col_widths=[7.5, 8.5],
)

spacer()
heading2("The platform in one sentence")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
p.paragraph_format.left_indent = Cm(0.8)
run = p.add_run(
    "Every morning by 07:30, every person on the team knows exactly what changed, "
    "what is at risk, what needs attention today, and where the opportunity is — "
    "without pulling a single report."
)
run.italic = True
run.font.size = Pt(11)
run.font.color.rgb = ORANGE

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════
# SECTION 3 — ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════════

heading1("3. Architecture")
body(
    "The platform is built on a layered architecture with a clear separation of concerns: "
    "data collection at the edges, AI processing in the runtime layer, structured storage in "
    "the data layer, and delivery through the portal and Teams interfaces."
)

spacer()
heading2("3.1 Architecture layers")
styled_table(
    ["Layer", "What it does", "Technology"],
    [
        (
            "User Access",
            "Authentication, role-based access control, session management",
            "Microsoft Entra ID SSO, Next.js middleware",
        ),
        (
            "Portal",
            "Structured modules for deep analysis, historical data, AI chat",
            "Next.js 16, TypeScript, React, hosted on AWS",
        ),
        (
            "Delivery",
            "Morning briefs to Teams channels, executive email to Denis",
            "Microsoft Graph API, Teams webhooks, Outlook",
        ),
        (
            "Data",
            "Persistent store for all agent outputs, metrics, and logs",
            "SQLite (Phase 1), PostgreSQL (Phase 2), AWS EC2",
        ),
        (
            "Runtime",
            "Cron-scheduled agent execution, model routing, output validation",
            "AWS EC2 t3.xlarge, Node.js, PM2",
        ),
        (
            "AI Models",
            "Three-tier model routing: router, default, escalation",
            "GPT-4o-mini, Claude Sonnet, Claude Opus",
        ),
        (
            "Agents",
            "Specialist agents that collect, analyse, and produce structured outputs",
            "5 Phase 1, 11 Phase 2",
        ),
        (
            "Integrations",
            "Live connections to all marketing and business data sources",
            "MCP connectors, REST APIs, Google SDKs",
        ),
        (
            "Reliability",
            "Backups, monitoring, secrets management, health checks",
            "AWS CloudWatch, S3, Secrets Manager",
        ),
    ],
    col_widths=[3.5, 6.5, 6],
)

spacer()
heading2("3.2 Three-tier AI model routing")
body(
    "Every agent uses a three-tier model routing design. This controls cost while ensuring "
    "quality where it matters. The router handles cheap classification decisions. Sonnet "
    "handles all production workloads. Opus is reserved for complex synthesis — primarily "
    "the executive brief and escalation cases."
)
spacer()
styled_table(
    ["Tier", "Model", "Purpose", "Cost per 1M tokens (input)"],
    [
        (
            "Router",
            "OpenAI GPT-4o-mini",
            "Intent classification, routing decisions, normalization",
            "$0.15",
        ),
        ("Default", "Claude Sonnet 4.6", "All 5 Phase 1 agents, standard daily execution", "$3.00"),
        (
            "Escalation",
            "Claude Opus 4.6",
            "Complex executive synthesis, high-stakes analysis only",
            "$5.00",
        ),
    ],
    col_widths=[2.5, 4, 7.5, 2],
)
spacer()
body(
    "Token controls: 50k tokens per run (Sonnet), 30k tokens per run (Opus), 200k tokens per agent per day. All model usage and cost logged automatically."
)

spacer()
heading2("3.3 Data flow")
info_box(
    "Morning intelligence cycle (daily, automated)",
    [
        "06:00  Performance Agent runs — collects GA4, Google Ads, Meta Ads data — writes to DB",
        "06:30  CRM & Lifecycle Agent runs — collects email, cart, churn data — writes to DB",
        "07:00  Reputation Agent runs — collects Trustpilot, Google Business data — writes to DB",
        "07:15  Marketing Ops Agent runs — checks tracking health, coupon dependency — writes to DB",
        "07:30  Cross-Channel Brief Agent reads all DB outputs — synthesises — posts to Teams + email",
        "        Portal reads precomputed DB outputs — never runs live AI on page load",
        "        AI Chat uses latest agent outputs as context for each conversation",
    ],
)

spacer()
heading2("3.4 Infrastructure")
styled_table(
    ["Component", "Phase 1 (current)", "Phase 2"],
    [
        (
            "Runtime host",
            "AWS EC2 t3.xlarge (2vCPU, 16GB RAM)",
            "Same host or t3.2xlarge depending on load",
        ),
        ("Database", "SQLite (portal.db + ananas_ai.db)", "PostgreSQL on EC2 with daily S3 backup"),
        ("Portal hosting", "EC2 behind nginx, HTTPS on IP", "ai.ananas.mk with SSL certificate"),
        (
            "Secrets",
            "Environment files on EC2 (restricted)",
            "AWS Secrets Manager for all credentials",
        ),
        ("Backups", "Manual S3 sync", "Nightly automated dump, EBS snapshots"),
        (
            "Monitoring",
            "PM2 process management, basic logs",
            "CloudWatch alarms, Teams health alerts",
        ),
        ("Auth", "Microsoft Entra ID SSO (live)", "Same, with expanded role matrix"),
    ],
    col_widths=[4, 6, 6],
)

spacer()
heading2("3.5 Role-based access")
body(
    "The portal enforces role-based access at the middleware layer. Every route is protected. Every user sees only what their role permits."
)
spacer()
styled_table(
    ["Role", "Access"],
    [
        ("Super Admin", "Full access: all modules, admin panel, user management, notifications"),
        ("Executive", "Executive summary, marketing overview, finance, cross-channel brief"),
        ("Marketing Head", "All marketing modules, admin panel, executive summary"),
        ("Performance Marketer", "Performance, CRM, reputation, marketing ops, overview"),
        ("CRM Specialist", "CRM & Lifecycle, performance overview, marketing ops"),
        ("Content & Brand", "Reputation, overview, influencer (Phase 2)"),
        ("Finance", "Finance module, commercial intelligence (Phase 2)"),
        ("HR", "HR module, employer branding (Phase 2)"),
    ],
    col_widths=[4.5, 11.5],
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════
# SECTION 4 — AGENTS
# ═══════════════════════════════════════════════════════════════════════════

heading1("4. Specialist Agents")
body(
    "Specialist agents are the core of the platform. Each agent owns one intelligence domain, "
    "runs on a defined schedule, connects to specific data sources, and produces a structured "
    "output that is validated, stored, and delivered. Agents do not share a runtime context "
    "and do not depend on each other except through the database."
)

spacer()
heading2("4.1 Phase 1 agents (live)")
styled_table(
    ["Agent", "Time", "Data sources", "Key metrics produced"],
    [
        (
            "Performance",
            "06:00",
            "GA4, Google Ads, Meta Ads, TikTok, LinkedIn, X",
            "POAS per campaign, blended ROAS, Shopping impression share, CPC trend, CVR by channel/device",
        ),
        (
            "CRM & Lifecycle",
            "06:30",
            "Email/CRM platform",
            "Cart abandonment rate, cart recovery rate, email revenue per send, churn rate 30/60/90d, repeat purchase",
        ),
        (
            "Reputation",
            "07:00",
            "Trustpilot API, Google Business Profile",
            "Rating, review count, response rate, sentiment trend, unanswered reviews, profile status",
        ),
        (
            "Marketing Ops",
            "07:15",
            "GA4, Orders API, Returns API",
            "Coupon dependency ratio, tracking health score, campaign analysis coverage, conversion discrepancy",
        ),
        (
            "Cross-Channel Brief",
            "07:30",
            "All 4 agent outputs (from DB)",
            "Daily team brief + executive summary — distributed to Teams and email",
        ),
    ],
    col_widths=[3.5, 1.5, 5.5, 5.5],
)

spacer()
heading2("4.2 Phase 2 agents (planned)")
styled_table(
    ["Agent", "Domain", "Schedule", "Business value"],
    [
        (
            "Customer Segmentation",
            "CRM",
            "Weekly (Mon)",
            "RFM segmentation, LTV by segment, churn risk scoring — reduces coupon dependency",
        ),
        (
            "Pricing Intelligence",
            "Commercial",
            "Daily",
            "Competitor price monitoring, repricing windows, margin-safe promotion detection",
        ),
        (
            "Search & Merchandising",
            "Conversion",
            "Daily",
            "Zero-result rate, failed search sessions, catalog gap identification",
        ),
        (
            "Listing Content Quality",
            "Catalog",
            "Weekly (Fri)",
            "250k product catalog quality score, thin listings ranked by GMV at risk",
        ),
        (
            "Competitor Intelligence",
            "Strategy",
            "Daily/Weekly",
            "Meta Ad Library monitoring, Google Auction Insights, competitor promotion detection",
        ),
        (
            "Organic & Merchandising",
            "SEO",
            "Daily",
            "Search Console, Ahrefs, product feed health, category indexation",
        ),
        (
            "Category Growth",
            "Marketplace",
            "Weekly (Mon)",
            "Category revenue by margin, return risk ranking, demand trend identification",
        ),
        (
            "Supplier Intelligence",
            "Marketplace",
            "Weekly (Tue)",
            "Supplier revenue contribution, co-marketing opportunities, campaign participation",
        ),
        (
            "Demand Forecasting",
            "Commercial",
            "Weekly (Wed)",
            "Demand spikes from search + category signals, seasonal pattern identification",
        ),
        (
            "Promo Simulator",
            "Commercial",
            "On-demand",
            "Pre-launch: expected GMV lift, margin impact, break-even volume per proposed discount",
        ),
        (
            "Product Feed",
            "Catalog",
            "Weekly (Thu)",
            "Google Shopping feed readiness, missing attributes, 250k+ product listing quality",
        ),
        (
            "Meeting Intelligence",
            "Operations",
            "On-upload",
            "Meeting transcripts to structured summaries, action items, Jira task creation",
        ),
        (
            "Knowledge Retrieval",
            "Operations",
            "On-demand",
            "Confluence + campaign archive search, institutional memory retrieval",
        ),
        (
            "Influencer & Partnership",
            "Marketing",
            "Weekly",
            "Creator ROI tracking, affiliate performance, co-marketing opportunity scoring",
        ),
        (
            "Traditional Media Correlation",
            "Marketing",
            "Weekly",
            "TV/OOH/Radio campaign dates correlated with branded search lift and direct traffic",
        ),
        (
            "Employer Branding",
            "HR",
            "Weekly",
            "LinkedIn employer presence, talent pipeline visibility, employer brand health",
        ),
    ],
    col_widths=[4, 2.5, 2.5, 7],
)

spacer()
heading2("4.3 Phase 3 agents (vision)")
styled_table(
    ["Agent", "Domain", "Description"],
    [
        (
            "Anomaly Detection",
            "Real-time",
            "Near-real-time monitoring for GMV drops, return rate spikes, payment failures, conversion collapse — requires event-triggered architecture",
        ),
        (
            "Attribution / MMM",
            "Strategy",
            "Marketing mix modelling and incrementality testing — measures true incremental revenue per channel, solves coupon attribution problem — requires 12+ months of clean tagged data",
        ),
    ],
    col_widths=[3.5, 2.5, 10],
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════
# SECTION 5 — PORTAL
# ═══════════════════════════════════════════════════════════════════════════

heading1("5. The Portal")
body(
    "The portal at ai.ananas.mk is the primary deep-analysis interface. It is a Next.js "
    "application secured by Microsoft Entra ID SSO, accessible only to authorised Ananas "
    "employees. Every module reads precomputed agent outputs from the database — the portal "
    "never triggers live AI calls on page load."
)

spacer()
heading2("5.1 Portal modules")
styled_table(
    ["Module", "Phase", "What it shows"],
    [
        (
            "Overview",
            "Live",
            "Daily cross-channel intelligence summary, agent status, last-run KPIs, revenue and sessions charts",
        ),
        (
            "Performance & Paid",
            "Live",
            "GA4 metrics, Google Ads and Meta Ads performance, channel breakdown, blended ROAS, date-range filter",
        ),
        (
            "CRM & Lifecycle",
            "Live",
            "Cart abandonment, recovery rate, email performance, retention KPIs, churn signals",
        ),
        (
            "Reputation",
            "Live",
            "Trustpilot and Google Business ratings, response rate, action checklist, unanswered reviews",
        ),
        (
            "Marketing Ops",
            "Live",
            "Tracking health, Search Console data, coupon dependency, alerts and ops notes",
        ),
        ("Influencers", "Live", "Influencer campaign tracking and creator performance"),
        ("Daily Brief", "Live", "Full cross-channel brief text, archive of past briefs"),
        (
            "AI Chat",
            "Live",
            "Conversational interface with full Ananas business context — asks questions, surfaces insights",
        ),
        (
            "Category Growth",
            "Phase 2",
            "Category revenue and margin ranking, return risk, demand trends",
        ),
        (
            "Promo Simulator",
            "Phase 2",
            "On-demand discount impact estimator: GMV lift, margin impact, break-even volume",
        ),
        (
            "Supplier Intelligence",
            "Phase 2",
            "Supplier revenue contribution, co-marketing opportunity scoring",
        ),
        (
            "Competitor Intelligence",
            "Phase 2",
            "Competitor ad activity, price positioning, auction impression share",
        ),
        (
            "Customer Segmentation",
            "Phase 2",
            "RFM segment distribution, churn risk, LTV by segment",
        ),
        (
            "Knowledge Search",
            "Phase 2",
            "Natural language search across Confluence, campaign archive, meeting notes",
        ),
        (
            "Finance",
            "Phase 2",
            "Finance-aligned KPI waterfall, contribution margin, budget tracking",
        ),
        ("HR", "Phase 2", "Team overview, employer branding metrics, Berry HR integration"),
    ],
    col_widths=[4, 2, 10],
)

spacer()
heading2("5.2 AI Chat")
body(
    "The AI Chat feature is live and available to authorised users. Each conversation is "
    "backed by a dynamic system prompt that is rebuilt on every request — loading the latest "
    "outputs from all five agents so the AI has real, current data to answer from."
)
spacer()
info_box(
    "What AI Chat knows (as of each conversation)",
    [
        "- GA4 revenue, sessions, users, and conversion rate from the last agent run",
        "- Google Ads and Meta Ads spend, ROAS, CPC, impressions, and clicks",
        "- Cart abandonment rate, email open rate, revenue per send, active subscribers",
        "- Trustpilot and Google Business ratings, response rate, claimed status",
        "- Tracking health, coupon dependency, Search Console data",
        "- Full text of the latest cross-channel brief",
        "- Full business context: team structure, known gaps, strategic priorities",
    ],
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════
# SECTION 6 — INTEGRATIONS
# ═══════════════════════════════════════════════════════════════════════════

heading1("6. Integrations")
body(
    "The platform connects to existing systems through the Model Context Protocol (MCP), "
    "REST APIs, and Google SDKs. All integrations in Phase 1 are read-only."
)
spacer()
styled_table(
    ["Integration", "Type", "Phase", "Status"],
    [
        ("Google Analytics 4 (GA4)", "Google SDK", "Phase 1", "Live and tested"),
        ("Google Ads (4 accounts)", "Google Ads API", "Phase 1", "Connected"),
        ("Google Search Console", "Search Console API", "Phase 1", "Connected"),
        ("Meta Ads (Facebook/Instagram)", "Meta Marketing API", "Phase 1", "Connected"),
        ("Trustpilot", "Trustpilot API", "Phase 1", "Connected"),
        ("Google Business Profile", "GBP API", "Phase 1", "Credentials needed"),
        ("Microsoft Teams (posting)", "Graph API", "Phase 1", "Connected"),
        ("Microsoft Outlook (email)", "Graph API", "Phase 1", "Mail.Send permission needed"),
        ("Microsoft Entra ID (SSO)", "OIDC / Graph", "Phase 1", "Live"),
        ("CRM / Email platform", "Platform API", "Phase 1", "Pending credentials"),
        ("TikTok Ads", "TikTok API", "Phase 1", "Pending credentials"),
        ("LinkedIn Ads", "LinkedIn API", "Phase 1", "Pending credentials"),
        ("X Ads", "X API", "Phase 1", "Account creation pending"),
        ("Orders API", "Internal REST", "Phase 1", "Pending backend team"),
        ("Returns API", "Internal REST", "Phase 1", "Pending backend team"),
        ("Product Catalog API", "Internal REST", "Phase 2", "Planned"),
        ("Supplier API", "Internal REST", "Phase 2", "Planned"),
        ("Margin / Finance API", "Internal REST", "Phase 2", "Planned"),
        ("Ahrefs / Semrush", "SEO API", "Phase 2", "Planned"),
        ("Jira", "Atlassian API", "Phase 2", "Planned"),
        ("Confluence", "Atlassian API", "Phase 2", "Planned"),
        ("Firebase / Adjust", "App analytics", "Phase 2", "Pending MK app launch"),
    ],
    col_widths=[5, 3.5, 2.5, 5],
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════
# SECTION 7 — PROJECT PLAN
# ═══════════════════════════════════════════════════════════════════════════

heading1("7. Project Plan")
body(
    "The project is structured across three phases. Phase 1 is operational. "
    "Phase 2 is the primary build period — full team rollout, all integrations, "
    "and the full agent suite. Phase 3 is the company-wide expansion."
)

spacer()
heading2("7.1 Phase 1 — Foundation (Months 1–3)")
info_box(
    "Status: Operational",
    [
        "The Phase 1 infrastructure is live. The portal runs at the server IP, five agents are scheduled,",
        "GA4 is integrated, Microsoft Entra ID SSO is working, and the first users are logged in.",
    ],
)
spacer()
styled_table(
    ["Milestone", "Status", "Details"],
    [
        ("AWS EC2 infrastructure", "Done", "t3.xlarge, nginx, PM2, HTTPS"),
        ("Portal — core shell", "Done", "Next.js 16, Entra ID SSO, role-based routing"),
        (
            "Portal — marketing modules",
            "Done",
            "Performance, CRM, Reputation, Ops, Overview, Brief",
        ),
        ("Portal — AI Chat", "Done", "Live with dynamic agent context"),
        ("Portal — Admin panel", "Done", "User management, invites, role changes"),
        ("Portal — Notifications", "Done", "System notifications, bell icon, read/unread"),
        ("GA4 integration", "Done", "Live — 464k sessions, 215k users, €13.4M revenue"),
        ("Google Ads integration", "Done", "4 accounts connected"),
        ("Meta Ads integration", "Done", "20 campaigns connected"),
        ("Performance Agent", "Done", "Daily cron 06:00"),
        ("CRM & Lifecycle Agent", "Done", "Daily cron 06:30"),
        ("Reputation Agent", "Done", "Daily cron 07:00"),
        ("Marketing Ops Agent", "Done", "Daily cron 07:15"),
        ("Cross-Channel Brief Agent", "Done", "Daily cron 07:30, Teams + email"),
        ("Output validation / guardrails", "Done", "All agent outputs validated before DB write"),
        ("TikTok / LinkedIn / X Ads", "Pending", "Credentials needed from platform accounts"),
        ("CRM / Email platform", "Pending", "Credentials needed from platform team"),
        ("Orders / Returns API", "Pending", "Access needed from backend team"),
        ("Mail.Send permission", "Pending", "Azure AD application permission needed"),
        ("Google Business Profile", "Pending", "GBP account ID and location ID needed"),
        ("ai.ananas.mk DNS", "Pending", "A record to 52.29.60.185 — IT action needed"),
        ("SSL certificate", "Pending", "After DNS — Let's Encrypt or company wildcard"),
    ],
    col_widths=[5.5, 2, 8.5],
)

spacer()
heading2("7.2 Phase 2 — Full Platform (Months 4–9)")
body(
    "Phase 2 delivers the complete platform: all integrations live, full agent suite, company-wide team rollout, and the portal accessible at ai.ananas.mk."
)
spacer()
styled_table(
    ["Work stream", "Month 4", "Month 5", "Month 6", "Month 7", "Month 8", "Month 9"],
    [
        ("DNS + SSL + production domain", "ai.ananas.mk live", "", "", "", "", ""),
        (
            "Remaining P1 integrations",
            "TikTok, LinkedIn, X, CRM, Orders/Returns",
            "Mail.Send, GBP",
            "",
            "",
            "",
            "",
        ),
        ("Team rollout (all 8 + mgmt)", "Invites sent, onboarding", "All active", "", "", "", ""),
        ("Customer Segmentation Agent", "Build + test", "Live", "", "", "", ""),
        ("Pricing Intelligence Agent", "", "Build + test", "Live", "", "", ""),
        ("Search & Merchandising Agent", "", "Build + test", "Live", "", "", ""),
        ("Listing Content Quality Agent", "", "", "Build + test", "Live", "", ""),
        ("Competitor Intelligence Agent", "", "", "Build + test", "Live", "", ""),
        ("Category Growth Agent", "", "", "", "Build + test", "Live", ""),
        ("Supplier Intelligence Agent", "", "", "", "Build + test", "Live", ""),
        ("Demand Forecasting Agent", "", "", "", "", "Build + test", "Live"),
        ("Promo Simulator Agent", "", "", "", "", "Build + test", "Live"),
        ("Product Feed Agent", "", "", "", "", "", "Build + test"),
        ("Meeting Intelligence Agent", "", "", "Build + test", "Live", "", ""),
        ("Knowledge Retrieval Agent", "", "", "", "Build + test", "Live", ""),
        ("Influencer & Partnership Agent", "", "", "", "", "Build + test", "Live"),
        ("Portal — new modules (P2)", "", "Start", "Running", "Running", "Complete", ""),
        ("PostgreSQL migration", "", "Plan", "Execute", "Validated", "", ""),
        ("AWS Secrets Manager", "Migrate", "", "", "", "", ""),
    ],
    col_widths=[5, 2.5, 2.5, 2.5, 2.5, 2.5, 2.5],
)

spacer()
heading2("7.3 Phase 3 — Company-wide expansion (Months 10–18)")
body(
    "Phase 3 expands the platform beyond marketing to cover the full business — commercial, finance, logistics, HR, and customer support. It also adds the most advanced capabilities: anomaly detection, marketing mix modelling, multilingual outputs, and creative automation."
)
spacer()
styled_table(
    ["Initiative", "Target department", "Months"],
    [
        ("Commercial intelligence (supplier, category, pricing)", "Commercial", "10-12"),
        ("Finance intelligence (margin waterfall, budget tracking)", "Finance", "11-13"),
        ("Customer Experience intelligence (CX, support, NPS)", "CX", "12-14"),
        ("Logistics intelligence (delivery performance, returns)", "Logistics", "13-15"),
        ("HR intelligence (Berry integration, employer branding)", "HR", "12-14"),
        ("Anomaly detection (real-time monitoring)", "All departments", "14-16"),
        ("Attribution / MMM (incrementality testing)", "Marketing", "16-18"),
        ("Multilingual outputs (Macedonian language)", "All", "15-17"),
        ("Video and creative automation layer", "Marketing / Brand", "16-18"),
        ("Self-learning content memory (marketing brain)", "Marketing", "17-18"),
    ],
    col_widths=[7, 4, 2.5],
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════
# SECTION 8 — BUDGET
# ═══════════════════════════════════════════════════════════════════════════

heading1("8. Budget")
body(
    "The platform budget covers three categories: infrastructure (cloud hosting and storage), "
    "AI model costs (API calls to Anthropic and OpenAI), and build investment (development "
    "time). The model is designed to be lean — a single EC2 host for Phase 1 and a carefully "
    "controlled token budget keep costs predictable."
)

spacer()
heading2("8.1 Phase 1 monthly operating costs (current)")
styled_table(
    ["Item", "Provider", "Monthly cost (EUR)", "Notes"],
    [
        ("EC2 t3.xlarge", "AWS", "~€110", "2 vCPU, 16GB RAM, runtime + portal + DB"),
        ("EBS storage (100GB)", "AWS", "~€10", "OS + application + database"),
        ("S3 backups", "AWS", "~€3", "Nightly DB dump, EBS snapshots"),
        ("CloudWatch", "AWS", "~€5", "Basic monitoring and alarms"),
        ("Data transfer", "AWS", "~€5", "Portal traffic, Teams posts, email"),
        ("Claude Sonnet (agents)", "Anthropic", "~€30-60", "5 agents x daily runs x ~50k tokens"),
        ("Claude Opus (escalation)", "Anthropic", "~€10-20", "Executive brief synthesis only"),
        ("GPT-4o-mini (routing)", "OpenAI", "~€2-5", "Classification and routing calls"),
        ("", "", "", ""),
        ("TOTAL Phase 1", "", "~€175-218/month", ""),
    ],
    col_widths=[5, 3, 3.5, 4.5],
)

spacer()
heading2("8.2 Phase 2 monthly operating costs (projected)")
styled_table(
    ["Item", "Monthly cost (EUR)", "Change from Phase 1"],
    [
        ("EC2 (same or t3.2xlarge)", "~€110-180", "Possible upgrade for 16-agent load"),
        ("Database (PostgreSQL + S3)", "~€15-25", "+€5-15 for S3 backup automation"),
        ("Claude Sonnet (16 agents)", "~€80-140", "3x agent volume, similar token discipline"),
        ("Claude Opus (escalation)", "~€20-40", "More complex synthesis across more briefs"),
        ("GPT-4o-mini (routing)", "~€5-10", "Proportional to agent volume"),
        ("Additional API costs (Ahrefs etc)", "~€50-100", "New paid data sources"),
        ("", "", ""),
        ("TOTAL Phase 2", "~€280-495/month", ""),
    ],
    col_widths=[6, 3.5, 6.5],
)

spacer()
heading2("8.3 Total investment summary (18 months)")
styled_table(
    ["Phase", "Duration", "Infra + AI costs", "Build investment", "Total"],
    [
        (
            "Phase 1 (Foundation)",
            "Months 1-3",
            "~€525-650",
            "Internal / already delivered",
            "~€525-650 + build",
        ),
        (
            "Phase 2 (Full Platform)",
            "Months 4-9",
            "~€1,680-2,970",
            "~€15,000-25,000 (agent build)",
            "~€17,000-28,000",
        ),
        (
            "Phase 3 (Company-wide)",
            "Months 10-18",
            "~€2,700-4,500",
            "~€25,000-40,000 (expansion)",
            "~€28,000-45,000",
        ),
        ("", "", "", "", ""),
        ("TOTAL (18 months)", "", "~€5,000-8,000", "~€40,000-65,000", "~€45,000-73,000"),
    ],
    col_widths=[4, 2.5, 3.5, 4.5, 3],
)
spacer()
info_box(
    "Budget notes",
    [
        "- Build investment assumes internal development with Claude Code assistance — no external agency",
        "- If external development is used, add €30,000-50,000 to build costs",
        "- AI model costs scale with agent volume and token discipline — caps are enforced per agent per day",
        "- Infrastructure costs can be reduced 30-40% by migrating to Hetzner (€60-100/month Phase 1)",
        "- Phase 3 timing and cost depend on business priorities and data availability (especially for MMM)",
    ],
)

spacer()
heading2("8.4 Cost per insight")
body(
    "With 16 agents producing structured intelligence across all major business domains, "
    "the all-in cost per actionable insight is estimated at under €0.50. A single avoided "
    "campaign mistake (wrong bid strategy, undetected Shopping gap, missed reputation crisis) "
    "typically recovers the monthly platform cost in full."
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════
# SECTION 9 — ROADMAP
# ═══════════════════════════════════════════════════════════════════════════

heading1("9. Roadmap")

heading2("9.1 Immediate priorities (next 30 days)")
styled_table(
    ["Priority", "Action", "Owner", "Dependency"],
    [
        (
            "1",
            "Grant Mail.Send application permission in Azure AD",
            "IT / Azure Admin",
            "Needed for invite emails",
        ),
        (
            "2",
            "Create ai.ananas.mk DNS A record → 52.29.60.185",
            "IT / DNS Admin",
            "Needed for production domain",
        ),
        (
            "3",
            "Provide Google Business Profile GBP_ACCOUNT_ID + LOCATION_ID",
            "Marketing",
            "Needed for reputation agent",
        ),
        ("4", "Provide TikTok Ads API credentials", "Marketing", "Needed for performance agent"),
        ("5", "Provide LinkedIn Ads API credentials", "Marketing", "Needed for performance agent"),
        (
            "6",
            "Create X (Twitter) Ads account and provide credentials",
            "Marketing",
            "Platform account needed first",
        ),
        ("7", "Connect CRM / email platform API", "CRM Specialist", "Needed for CRM agent"),
        ("8", "Grant Orders and Returns API read access", "Backend Team", "Needed for ops agent"),
        (
            "9",
            "Install SSL certificate on ai.ananas.mk after DNS",
            "Platform Team",
            "After DNS action",
        ),
        ("10", "Invite full marketing team (8 people) to portal", "Admin (Zharko)", "Ready now"),
    ],
    col_widths=[1, 6.5, 3.5, 4.5],
)

spacer()
heading2("9.2 Q2 2026 (Months 4–6)")
bullet("ai.ananas.mk live with SSL certificate and production domain")
bullet("Full team using the portal daily — all 8 marketing team members + Denis")
bullet("Customer Segmentation Agent live — RFM segments, churn risk scoring")
bullet("Pricing Intelligence Agent live — competitor price monitoring")
bullet("Search & Merchandising Agent live — zero-result rate, catalog gap detection")
bullet("Meeting Intelligence Agent live — transcripts to Jira tasks")
bullet("PostgreSQL migration completed — from SQLite to production database")
bullet("AWS Secrets Manager live — all credentials migrated")

spacer()
heading2("9.3 Q3 2026 (Months 7–9)")
bullet("Category Growth Agent live — marketplace revenue and margin ranking")
bullet("Supplier Intelligence Agent live — co-marketing opportunity scoring")
bullet("Demand Forecasting Agent live — seasonal demand spike detection")
bullet("Promo Simulator Agent live — pre-launch discount impact estimation")
bullet("Knowledge Retrieval Agent live — Confluence + campaign archive search")
bullet("Influencer & Partnership Agent live — creator ROI tracking")
bullet("All portal Phase 2 modules live and accessible")
bullet("Teams bot v1 live — promo simulation and knowledge retrieval via chat")

spacer()
heading2("9.4 Q4 2026 and beyond (Months 10–18)")
bullet("Company-wide expansion: Commercial, Finance, CX, Logistics, HR modules")
bullet("Anomaly Detection Agent — near-real-time GMV and conversion monitoring")
bullet("Attribution / MMM Agent — true incremental revenue measurement per channel")
bullet("Multilingual outputs — Macedonian language for all briefs and portal content")
bullet("Creative automation layer — AI-generated ad copy and brief drafts")
bullet("Self-learning content memory — the platform learns from its own history")

spacer()
heading2("9.5 Strategic milestones")
styled_table(
    ["Milestone", "Target date", "What it means"],
    [
        (
            "Platform operational",
            "March 2026",
            "5 agents running, portal live, first users — ACHIEVED",
        ),
        (
            "Full team on platform",
            "April 2026",
            "All 8 marketing team members + Denis using portal daily",
        ),
        (
            "All Phase 1 integrations live",
            "April 2026",
            "TikTok, LinkedIn, X, CRM, Orders, email all connected",
        ),
        ("Production domain live", "April 2026", "ai.ananas.mk with SSL certificate"),
        (
            "Phase 2 agents live",
            "September 2026",
            "16 specialist agents covering all major intelligence domains",
        ),
        (
            "Company-wide launch",
            "Q4 2026",
            "Commercial, Finance, HR, CX departments using the platform",
        ),
        (
            "Full attribution model live",
            "Q1 2027",
            "Marketing mix model measuring true incremental revenue",
        ),
        (
            "Platform self-improving",
            "Q2 2027",
            "Self-learning content memory, multilingual outputs",
        ),
    ],
    col_widths=[5, 3, 8],
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════
# SECTION 10 — PENDING ACTIONS
# ═══════════════════════════════════════════════════════════════════════════

heading1("10. Pending Actions and Dependencies")
body(
    "The following items are blocking full Phase 1 completion. They require action "
    "from IT, Azure administration, or platform account owners."
)
spacer()
styled_table(
    ["Ticket", "Description", "Who", "Impact if delayed"],
    [
        (
            "IT-001",
            "Grant Mail.Send APPLICATION permission in Azure AD (app: Ananas AI Portal). Required for automated invite emails and system notifications.",
            "IT / Azure Admin",
            "Invite emails fail. Manual role assignment only.",
        ),
        (
            "IT-002",
            "Grant User.Read.All APPLICATION permission in Azure AD. Required for HR user sync and org chart features in Phase 2.",
            "IT / Azure Admin",
            "HR sync not possible in Phase 2.",
        ),
        (
            "IT-003",
            "Create DNS A record: ai.ananas.mk → 52.29.60.185. Required for production domain access.",
            "IT / DNS Admin",
            "Portal accessible by IP only. No clean URL for team.",
        ),
        (
            "IT-004",
            "Issue or install SSL certificate for ai.ananas.mk. Let's Encrypt or company wildcard certificate.",
            "IT / Platform Team",
            "Portal not accessible on production domain.",
        ),
        (
            "MKTG-001",
            "Provide Google Business Profile GBP_ACCOUNT_ID and GBP_LOCATION_ID for reputation agent.",
            "Marketing Team",
            "Google Business data missing from Reputation Agent.",
        ),
        (
            "MKTG-002",
            "Provide TikTok Ads API credentials (App ID, App Secret, Access Token, Advertiser ID).",
            "Marketing Team",
            "TikTok excluded from Performance Agent.",
        ),
        (
            "MKTG-003",
            "Provide LinkedIn Ads API credentials (Client ID, Client Secret, Account ID).",
            "Marketing Team",
            "LinkedIn excluded from Performance Agent.",
        ),
        (
            "MKTG-004",
            "Create X (Twitter) Ads account and provide API credentials.",
            "Marketing Team",
            "X Ads excluded from Performance Agent.",
        ),
        (
            "MKTG-005",
            "Provide CRM / email platform API credentials (API key, account ID, base URL).",
            "CRM Specialist",
            "CRM & Lifecycle Agent has no data source.",
        ),
        (
            "DEV-001",
            "Grant read-only access to Orders API and Returns API for marketing ops agent.",
            "Backend Team",
            "Marketing Ops Agent cannot calculate coupon dependency or margin waterfall.",
        ),
        (
            "DEV-002",
            "Grant read-only access to Product Catalog API (250k products). Required for Phase 2 product feed and listing quality agents.",
            "Backend Team",
            "Product Feed and Listing Quality agents blocked.",
        ),
    ],
    col_widths=[2, 6.5, 3, 4.5],
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════
# SECTION 11 — GOVERNANCE
# ═══════════════════════════════════════════════════════════════════════════

heading1("11. Governance and Reliability")

heading2("11.1 Output quality")
bullet(
    "Every agent output is validated against a JSON schema before being written to the database."
)
bullet(
    "A missing or invalid run_type field is always a hard failure — nothing is written or posted."
)
bullet("Validation failures are logged and trigger a Teams alert. They are never silently ignored.")
bullet("Sample-mode outputs (run_type: sample) are labelled [SAMPLE DATA] in all Teams messages.")
bullet(
    "Switching any agent from sample mode to live requires a manual confirmation step — never automated."
)

spacer()
heading2("11.2 Security")
bullet(
    "All credentials are stored in environment files on the EC2 host with restricted permissions. Phase 2 migrates to AWS Secrets Manager."
)
bullet("No credentials are stored in the GitHub repository or in any portal frontend code.")
bullet("The portal enforces role-based access at the middleware layer on every request.")
bullet("All integrations in Phase 1 are read-only. No writes to any business system.")
bullet("Microsoft Entra ID SSO is the only authentication method — no local passwords.")

spacer()
heading2("11.3 Backups and recovery")
bullet("Daily SQLite database dumps are stored on the EC2 host. Phase 2 adds automated S3 backup.")
bullet("The portal shows the last successful agent output with a timestamp if a fresh run fails.")
bullet("PM2 automatically restarts the portal process if it crashes.")
bullet("CloudWatch monitors EC2 health and triggers alerts on anomalies.")

spacer()
heading2("11.4 Repository and documentation")
body(
    "All platform code is in the private GitHub repository (github.com/zhapostolski/ananas-ai). Every architectural decision is documented as an Architecture Decision Record (ADR) in docs/decisions/. Significant changes require updates to the architecture document, agents.json, and CHANGELOG."
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════
# APPENDIX — METRIC COVERAGE
# ═══════════════════════════════════════════════════════════════════════════

heading1("Appendix — Metric Coverage")
styled_table(
    ["Category", "Phase 1 coverage", "Phase 2 coverage"],
    [
        (
            "Paid acquisition (all channels)",
            "Full — Google, Meta, TikTok, LinkedIn, X",
            "Same + competitor benchmarking",
        ),
        ("GA4 analytics", "Full — live and tested", "Same + app analytics (Firebase)"),
        (
            "Google Shopping",
            "Shopping Impression Share tracked",
            "Full feed quality + 250k product coverage",
        ),
        (
            "CRM / lifecycle",
            "Full — cart recovery, email, churn",
            "RFM segmentation, LTV by segment",
        ),
        (
            "Reputation",
            "Full — Trustpilot + Google Business",
            "Sentiment trend, response automation",
        ),
        ("Coupon dependency", "Full — critical monitoring metric", "Attribution / MMM layer"),
        ("POAS per campaign", "Full — campaign-level, not just blended", "Same"),
        (
            "Organic / SEO",
            "Search Console + Ahrefs (Phase 1)",
            "Full SEO automation + keyword tracking",
        ),
        (
            "Category-level profitability",
            "Partial — depends on business API access",
            "Full — Category Growth Agent",
        ),
        ("Product feed health", "Not covered", "Full — 250k catalog quality scoring"),
        ("Supplier intelligence", "Not covered", "Full — Supplier Intelligence Agent"),
        ("Demand forecasting", "Not covered", "Full — Demand Forecasting Agent"),
        ("Competitor intelligence", "Not covered", "Meta Ad Library + Auction Insights"),
        ("Pricing intelligence", "Not covered", "Full — Pricing Intelligence Agent"),
        ("Meeting intelligence", "Not covered", "Transcripts to Jira tasks"),
        ("Institutional memory", "Not covered", "Confluence + campaign archive search"),
        ("Traditional media lift", "Not covered", "TV/OOH/Radio correlation"),
        ("App metrics", "Not covered", "Pending MK app launch"),
        ("Attribution / MMM", "Not covered", "Phase 3 — requires 12+ months clean data"),
    ],
    col_widths=[4.5, 5.5, 6],
)

spacer()

# Footer note
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(20)
run = p.add_run("Ananas AI Platform — Internal Document — Version 1.0 — March 2026")
run.font.size = Pt(9)
run.font.color.rgb = GRAY
run.italic = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ── Save ───────────────────────────────────────────────────────────────────
output_path = "/home/zapostolski/projects/ananas-ai/docs/Ananas-AI-Platform-Master.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
